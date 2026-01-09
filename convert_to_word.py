"""
تحويل ملف Markdown إلى Word مع تنسيق عربي صحيح
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_rtl(paragraph):
    """تعيين اتجاه النص من اليمين لليسار"""
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def set_arabic_font(run):
    """تعيين خط عربي مناسب"""
    run.font.name = 'Traditional Arabic'
    run._element.rPr.rFonts.set(qn('w:cs'), 'Traditional Arabic')
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')

# إنشاء مستند Word جديد
doc = Document()
doc.core_properties.title = 'دليل المستخدم الشامل - برنامج المكتب الانتخابي'
doc.core_properties.author = 'المكتب الانتخابي'

# تنسيق الأنماط
for style_name in ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Normal']:
    if style_name in doc.styles:
        style = doc.styles[style_name]
        style.font.name = 'Traditional Arabic'
        if hasattr(style, 'paragraph_format'):
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# قراءة ملف Markdown
with open(r'C:\Users\2025\.gemini\antigravity\scratch\electoral_office\دليل_المستخدم_الشامل.md', 'r', encoding='utf-8') as f:
    content = f.read()

lines = content.split('\n')
in_code_block = False

for line in lines:
    original_line = line
    line = line.strip()
    
    # تجاهل الكود بلوكس
    if line.startswith('```'):
        in_code_block = not in_code_block
        continue
    
    if in_code_block:
        p = doc.add_paragraph(original_line)
        p.style = 'No Spacing'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # الكود من اليسار
        if p.runs:
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
        continue
    
    # معالجة العناوين
    if line.startswith('# ') and not line.startswith('## '):
        heading_text = line[2:]
        if heading_text.strip():
            p = doc.add_heading(heading_text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl(p)
            for run in p.runs:
                set_arabic_font(run)
                run.font.size = Pt(20)
                run.font.color.rgb = RGBColor(0, 51, 102)
    elif line.startswith('#### '):
        heading_text = line[5:]
        if heading_text.strip():
            p = doc.add_heading(heading_text, level=4)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl(p)
            for run in p.runs:
                set_arabic_font(run)
    elif line.startswith('### '):
        heading_text = line[4:]
        if heading_text.strip():
            p = doc.add_heading(heading_text, level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl(p)
            for run in p.runs:
                set_arabic_font(run)
    elif line.startswith('## '):
        heading_text = line[3:]
        if heading_text.strip():
            p = doc.add_heading(heading_text, level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl(p)
            for run in p.runs:
                set_arabic_font(run)
                run.font.size = Pt(16)
    # القوائم
    elif line.startswith('- ') or line.startswith('* '):
        text = line[2:]
        if text.strip():
            # إزالة رموز التنسيق
            text = text.replace('**', '').replace('__', '')
            p = doc.add_paragraph(text, style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl(p)
            for run in p.runs:
                set_arabic_font(run)
    # الاقتباسات والملاحظات
    elif line.startswith('> '):
        text = line[2:]
        if text.strip():
            text = text.replace('**', '').replace('__', '')
            p = doc.add_paragraph(text)
            p.style = 'Quote'
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl(p)
            for run in p.runs:
                set_arabic_font(run)
                run.font.color.rgb = RGBColor(102, 102, 102)
    # الخطوط الأفقية
    elif line.startswith('---'):
        p = doc.add_paragraph()
        # فاصل بسيط
    # النصوص العادية
    elif line.strip():
        # إزالة رموز Markdown
        text = line.replace('**', '').replace('__', '').replace('`', '')
        if text.strip() and not text.startswith('[') and not text.startswith('http'):
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl(p)
            for run in p.runs:
                set_arabic_font(run)
                run.font.size = Pt(12)

# إضافة ترقيم الصفحات في التذييل
section = doc.sections[0]
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.text = "صفحة "
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_para.add_run()
from docx.oxml import parse_xml
fldChar1 = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
instrText = parse_xml(r'<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE </w:instrText>')
fldChar2 = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
run._r.append(fldChar1)
run._r.append(instrText)
run._r.append(fldChar2)

# حفظ المستند
output_path = r'C:\Users\2025\.gemini\antigravity\scratch\electoral_office\دليل_المستخدم_الشامل.docx'
doc.save(output_path)
print(f'✅ تم إنشاء ملف Word بتنسيق عربي صحيح في: {output_path}')
print(f'📊 الحجم: {len(lines)} سطر')
